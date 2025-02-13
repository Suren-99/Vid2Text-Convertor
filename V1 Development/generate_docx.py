from docx import Document
import os


def create_docx(text, folder_path, filename, heading='Generated Text'):
    """Creates a DOCX file with the provided text."""
    filename += '.docx'
    file_path = os.path.join(folder_path, filename)

    try:
        doc = Document()
        if heading:
            doc.add_heading(heading, 0)
        doc.add_paragraph(text)
        doc.save(file_path)
        return file_path
    except IOError as e:
        print(f"IOError: {e}. Could not create the DOCX file.")
    except Exception as e:
        print(f"Unexpected error: {e}")
